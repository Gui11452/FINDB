import tempfile
from docx import Document
from docx.shared import Inches

def replace_placeholder(doc, placeholder, replacement):
    # print(placeholder, replacement)
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                print(placeholder, item.text, placeholder == item.text)
                if placeholder in item.text:
                    item.text = item.text.replace(placeholder, replacement)
                    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder(cell, placeholder, replacement)

def insert_image(doc, placeholder, image_path, width=None, height=None):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if placeholder in item.text:
                    item.clear()  # Clear the placeholder text
                    run = paragraph.add_run()
                    if width and height:
                        run.add_picture(image_path, width=width, height=height)  # Insert the image with specified size
                    elif width:
                        run.add_picture(image_path, width=width)  # Insert the image with specified width
                    elif height:
                        run.add_picture(image_path, height=height)  # Insert the image with specified height
                    else:
                        run.add_picture(image_path)  # Insert the image with default size
                    return  # Assuming one replacement per placeholder

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                insert_image(cell, placeholder, image_path, width, height)

def create_confirmation_ticket(template_path, replacements, images):
    # Abre o documento existente
    doc = Document(template_path)
    
    # Substitui os placeholders pelas variáveis
    for placeholder, replacement in replacements.items():
        replace_placeholder(doc, placeholder, replacement)
    
    # Insere as imagens nos placeholders
    for placeholder, image_info in images.items():
        image_path = image_info['path']
        width = image_info.get('width', None)
        height = image_info.get('height', None)
        insert_image(doc, placeholder, image_path, width, height)
    
    # Cria um arquivo temporário
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.seek(0)
    return temp_file


